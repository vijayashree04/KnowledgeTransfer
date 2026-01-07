import streamlit as st
import os
from dotenv import load_dotenv
from datetime import datetime, timezone, timedelta

# Load environment variables first
load_dotenv()

def format_date_ist(date_string):
    """Format ISO date string to DD-MM-YYYY HH:MM AM/PM in IST."""
    if not date_string or date_string == 'N/A':
        return 'N/A'
    try:
        # Parse the ISO format date
        if isinstance(date_string, str):
            # Handle different ISO formats
            if '+' in date_string or date_string.endswith('Z'):
                # Remove timezone info and parse
                date_str_clean = date_string.replace('+00:00', '').replace('Z', '')
                if '.' in date_str_clean:
                    # Has microseconds
                    dt = datetime.fromisoformat(date_str_clean)
                else:
                    dt = datetime.fromisoformat(date_str_clean)
                # Assume UTC if no timezone info
                dt = dt.replace(tzinfo=timezone.utc)
            else:
                dt = datetime.fromisoformat(date_string)
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
        else:
            dt = date_string
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
        
        # Convert to IST (UTC+5:30)
        ist_offset = timedelta(hours=5, minutes=30)
        ist_tz = timezone(ist_offset)
        dt_ist = dt.astimezone(ist_tz)
        
        # Format as DD-MM-YYYY HH:MM AM/PM
        formatted = dt_ist.strftime('%d-%m-%Y %I:%M %p')
        return formatted
    except Exception as e:
        # If parsing fails, return original string
        return str(date_string)

# Now import modules that depend on environment variables
import auth
import document_store
import gemini_utils
import team_store
import user_store
import landing_page
import theme

# Page Config
st.set_page_config(
    page_title="ClearKT",
    page_icon="assets/clearkt_logo.png",
    layout="wide",
    initial_sidebar_state="auto"
)
st.title("ClearKT")
# Apply global theme
theme.apply_global_theme()

# Navigation logic - use session state
if "current_page" not in st.session_state:
    st.session_state.current_page = "landing"

# Check authentication
is_authenticated = auth.check_auth()

# Route to appropriate page
if not is_authenticated:
    if st.session_state.current_page == "signup":
        auth.signup_page()
    elif st.session_state.current_page == "team_choice":
        auth.team_choice_page()
    elif st.session_state.current_page == "create_team":
        auth.create_team_page()
    elif st.session_state.current_page == "join_team":
        auth.join_team_page()
    elif st.session_state.current_page == "login":
        auth.login_page()
    else:
        landing_page.show_landing_page()
else:
    
    st.markdown("""
        <style>
        [data-testid="column"] {
            display: flex;
            align-items: center;
        }
        /* Override theme for ClearKT header - Force colors */
        .clearkt-clear-span,
        span.clearkt-clear-span,
        .clearkt-header-clear,
        span.clearkt-header-clear {
            color: #3b82f6 !important;
            display: inline !important;
        }
        .clearkt-kt-span,
        span.clearkt-kt-span,
        .clearkt-header-kt,
        span.clearkt-header-kt {
            color: #000000 !important;
            display: inline !important;
        }
        /* Override any theme color rules for these spans */
        h1 .clearkt-clear-span,
        h1 span.clearkt-clear-span {
            color: #3b82f6 !important;
        }
        h1 .clearkt-kt-span,
        h1 span.clearkt-kt-span {
            color: #000000 !important;
        }
        /* Force ClearKT header title size - must override all other rules */
        .clearkt-main-header h1,
        h1.clearkt-main-header,
        .title_col h1,
        div[data-testid="column"]:has(.clearkt-main-header) h1,
        .stMarkdown:has(.clearkt-main-header) h1 {
            font-size: 75px !important;
            margin: 0 !important;
            font-weight: 1000 !important;
            letter-spacing: -0.02em !important;
            font-family: "Inter", -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
            line-height: 1 !important;
        }
        </style>
        """, unsafe_allow_html=True)
    # Logout button in the main area (top right)
    col_title, col_logout = st.columns([4, 1])
    with col_title:
        # Logo + title in a row with better alignment
        logo_col, title_col = st.columns([1, 5])
        
        with logo_col:
            st.image("assets/clearkt_logo.png", width=175)  # Larger, clearer logo

        with title_col:
            st.markdown(
                """
                <style>
                .clearkt-clear-span {
                    color: #3b82f6 !important;
                    display: inline !important;
                }
                .clearkt-kt-span {
                    color: #000000 !important;
                    display: inline !important;
                }
                </style>
                <div class="clearkt-main-header" style='display: flex; align-items: center; padding-top: 2rem;'>
                    <h1 class="clearkt-main-header" style='margin:0 !important; font-size: 75px !important; font-weight: 1000 !important; letter-spacing: -0.02em !important; font-family: "Inter", -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important; line-height: 1 !important;'>
                        <span class="clearkt-clear-span">Clear</span><span class="clearkt-kt-span">KT</span>
                    </h1>
                </div>
                """,
                unsafe_allow_html=True
            )
            
        if "team_name" in st.session_state:
            st.caption(f"Team: {st.session_state.team_name}")
    with col_logout:
        st.markdown("<br>", unsafe_allow_html=True)  # Add some spacing
        if st.button("Logout", type="secondary", use_container_width=True):
            auth.logout()
    
    # Welcome message banner (below logo, above tabs) - auto-hide after 3 minutes
    import time
    user_name = st.session_state.get("name", st.session_state.get("email", "User"))
    
    # Initialize login time if not set
    if "login_time" not in st.session_state:
        st.session_state.login_time = time.time()
    
    # Check if 3 minutes (180 seconds) have passed
    elapsed_time = time.time() - st.session_state.login_time
    show_welcome = elapsed_time < 180  # 3 minutes
    
    if show_welcome:
        st.markdown(f"""
        <div style='background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); padding: 1.5rem; border-radius: 12px; margin-bottom: 2rem; border-left: 4px solid #3b82f6;'>
            <h2 style='color: #1e293b; margin: 0 0 0.5rem 0; font-size: 1.75rem; font-weight: 600;'>
                👋 Welcome back, {user_name}!
            </h2>
            <p style='color: #475569; margin: 0; font-size: 1.1rem;'>
                Ready to explore your team's knowledge base? Upload documents, view summaries, or chat with the AI assistant.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Check if user is team lead
    import team_store
    team_id = st.session_state.get("team_id")
    user_email = st.session_state.get("email", "")
    is_lead = st.session_state.get("is_team_lead", False)
    if not is_lead and team_id:
        is_lead = team_store.is_team_lead(team_id, user_email)
    
    # Tabs - add Team Settings if user is team lead
    if is_lead:
        tab1, tab2, tab3, tab4 = st.tabs(["📤 Upload Documents", "📝 Document Summaries", "🤖 KT Chatbot", "⚙️ Team Settings"])
    else:
        tab1, tab2, tab3 = st.tabs(["📤 Upload Documents", "📝 Document Summaries", "🤖 KT Chatbot"])

    # Tab 1: Upload
    with tab1:
        # Header section
        st.markdown("""
        <div style='margin-bottom: 1rem; text-align: center;'>
            <p style='color: #1e293b; margin: 1; font-size: 1.5rem; font-weight: 500; line-height: 1.5;'>Add documents to your team's knowledge base</p>
        </div>
        """, unsafe_allow_html=True)
        
        # File uploader with hidden label
        uploaded_file = st.file_uploader(
            "Upload document",
            type=["txt", "md", "py", "js", "json", "pdf", "docx", "doc"],
            label_visibility="collapsed",
            
        )
        
        # Show file info and format info
        
        # Show file info and process button if file is uploaded
        if uploaded_file is not None:
            st.markdown("<br>", unsafe_allow_html=True)


            
            # Process and Upload button - centered
            col1, col2, col3 = st.columns([1, 2.5, 1])
            with col2:
                if st.button("Upload & Summarize", type="primary", use_container_width=True):
                    with st.spinner("Uploading and Summarizing..."):

                        # Save File
                        team_id = st.session_state.get("team_id")
                        user_email = st.session_state.get("email", "unknown")
                        team_name = st.session_state.get("team_name")

                        metadata = document_store.save_uploaded_file(
                            uploaded_file,
                            user_email,
                            team_id,
                            team_name
                        )

                        # Success Message (Centered, Consistent Width)
                        st.markdown("""
                        <div style="display:flex; justify-content:center; margin:2rem 0;">
                            <div style="
                                background:#f0fdf4;
                                padding:0.75rem 1.5rem;
                                border-radius:8px;
                                border-left:4px solid #22c55e;
                                max-width:750px;
                                width:100%;
                            ">
                                <p style="margin:0; color:#166534; font-weight:500;">
                                    Document uploaded successfully
                                </p>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        # Generate Summary
                        summary = gemini_utils.summarize_document_short(
                            metadata["content"],
                            uploaded_file.name
                        )

                        document_store.update_document_summary(
                            uploaded_file.name,
                            summary,
                            team_id
                        )

                        # Summary Heading (Aligned with Content)
                        st.markdown("""
                        <div style="display:flex; justify-content:center; margin:2rem 0 1rem 0;">
                            <h3 style="
                                color:#1e293b;
                                max-width:750px;
                                width:100%;
                                text-align:left;
                            ">
                                📋 Document Summary
                            </h3>
                        </div>
                        """, unsafe_allow_html=True)

                        if summary.strip():
                            # Format summary as bullet points
                            lines = summary.strip().split("\n")
                            bullet_lines = []

                            for line in lines:
                                line = line.strip()
                                if line:
                                    if line.startswith(("-", "*", "·", "•")):
                                        line = "• " + line.lstrip("-*·• ").strip()
                                    else:
                                        line = f"• {line}"
                                    bullet_lines.append(line)

                            # Limit to max 6 bullets
                            bullet_lines = bullet_lines[:6]

                            # Summary Card
                            st.markdown(f"""
                            <div style="display:flex; justify-content:center; margin-bottom:2rem;">
                                <div style="
                                    background:#ffffff;
                                    padding:2rem;
                                    border-radius:12px;
                                    box-shadow:0 1px 3px rgba(0,0,0,0.1);
                                    max-width:750px;
                                    width:100%;
                                    text-align:left;
                                    line-height:1.6;
                                    color:#0f172a;
                                ">
                                    {'<br>'.join(bullet_lines)}
                                </div>
                            </div>
                            """, unsafe_allow_html=True)

                        else:
                            st.info("No summary available.")


    # Tab 2: Summaries
    with tab2:
        st.header("Document Repository")
        docs = document_store.get_documents(team_id)
        
        
        if not docs:
            st.info("No documents uploaded yet for your team.")
        else:
            st.info(f"Showing {len(docs)} document(s) for your team.")
            for doc in docs:
                col1, col2 = st.columns([4, 1])
                
                with col1:
                    upload_date_formatted = format_date_ist(doc.get('upload_date', ''))
                    with st.expander(f"📄 {doc['filename']} (Uploaded by {doc['uploaded_by']} on {upload_date_formatted})"):
                        st.markdown("### Detailed Summary")
                        
                        # Check if detailed summary exists in Supabase
                        detailed_summary = doc.get("detailed_summary")
                        
                        # If not in database, generate and save it
                        if not detailed_summary:
                            with st.spinner("Generating detailed summary..."):
                                # Get full document content
                                full_content = doc.get("content", "")
                                
                                # If content is not available, try to get from database
                                if not full_content or len(full_content) < 100:
                                    try:
                                        docs = document_store.get_documents(team_id)
                                        for d in docs:
                                            if d.get('filename') == doc['filename']:
                                                full_content = d.get("content", "")
                                                if full_content:
                                                    break
                                    except:
                                        pass
                                
                                if full_content and len(full_content) > 50:
                                    # Generate detailed summary
                                    detailed_summary = gemini_utils.summarize_document_detailed(full_content, doc['filename'])
                                    
                                    # Save to Supabase
                                    try:
                                        document_store.update_document_detailed_summary(doc['filename'], detailed_summary, team_id)
                                        # Update the doc dict so it's available for next time
                                        doc['detailed_summary'] = detailed_summary
                                    except Exception as e:
                                        st.warning(f"Summary generated but could not be saved to database: {str(e)}")
                                else:
                                    detailed_summary = "⚠️ Content not available for detailed summary generation. The document content may not have been extracted during upload."
                        
                        # Display detailed summary with markdown formatting
                        if detailed_summary:
                            # Process markdown to convert standalone **text** to smaller headings
                            import re
                            # Convert standalone **text** lines to smaller markdown headings (#### = h4)
                            lines = detailed_summary.split('\n')
                            processed_lines = []
                            for line in lines:
                                # Check if line is a standalone bold heading (starts and ends with **)
                                if re.match(r'^\s*\*\*.*\*\*\s*$', line):
                                    # Convert to smaller heading (#### = h4, which is smaller than h2/h3)
                                    heading_text = re.sub(r'\*\*', '', line).strip()
                                    processed_lines.append(f"#### {heading_text}")
                                else:
                                    processed_lines.append(line)
                            processed_summary = '\n'.join(processed_lines)
                            
                            # Add custom CSS to make all headings in summary smaller (but not the header)
                            st.markdown("""
                                <style>
                                /* Target headings in the summary section only - exclude header */
                                .stMarkdown:not(:has(.clearkt-main-header)) h1 { font-size: 1.3em !important; }
                                .stMarkdown:not(:has(.clearkt-main-header)) h2 { font-size: 1.2em !important; }
                                .stMarkdown:not(:has(.clearkt-main-header)) h3 { font-size: 1.15em !important; }
                                .stMarkdown:not(:has(.clearkt-main-header)) h4 { font-size: 1.1em !important; font-weight: 600 !important; }
                                .stMarkdown:not(:has(.clearkt-main-header)) h5 { font-size: 1.05em !important; }
                                .stMarkdown:not(:has(.clearkt-main-header)) h6 { font-size: 1em !important; }
                                /* Ensure header title stays large */
                                .clearkt-main-header h1,
                                h1.clearkt-main-header {
                                    font-size: 96px !important;
                                }
                                /* Override code block styling - blue background with white text */
                                .stMarkdown code,
                                .stMarkdown pre,
                                .stMarkdown pre code {
                                    background-color: #3b82f6 !important;
                                    color: #ffffff !important;
                                    border-radius: 4px !important;
                                    padding: 0.2em 0.4em !important;
                                }
                                .stMarkdown pre {
                                    background-color: #3b82f6 !important;
                                    color: #ffffff !important;
                                    border: 1px solid #2563eb !important;
                                    padding: 1rem !important;
                                }
                                .stMarkdown pre code {
                                    background-color: transparent !important;
                                    color: #ffffff !important;
                                    padding: 0 !important;
                                }
                                </style>
                            """, unsafe_allow_html=True)
                            
                            # Display the processed summary
                            st.markdown(processed_summary)
                            
                            # Download options
                            st.markdown("---")
                            st.markdown("### Download Summary")
                            
                            col_dl1, col_dl2, col_dl3 = st.columns(3)
                            
                            with col_dl1:
                                # Download as TXT
                                summary_txt = detailed_summary.encode('utf-8')
                                st.download_button(
                                    label="📄 Download as TXT",
                                    data=summary_txt,
                                    file_name=f"{doc['filename']}_summary.txt",
                                    mime="text/plain",
                                    key=f"download_txt_{doc['filename']}"
                                )
                            
                            
                            with col_dl3:
                                # Download as DOCX
                                try:
                                    from docx import Document
                                    from io import BytesIO
                                    import re
                                    
                                    # Helper function to add text with bold formatting
                                    def add_formatted_text(paragraph, text):
                                        """Adds text to paragraph, handling **bold** and `code` markdown."""
                                        # Split by ** and ` to find formatted sections
                                        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
                                        for part in parts:
                                            if part.startswith('**') and part.endswith('**'):
                                                # Bold text - remove **
                                                bold_text = part.replace('**', '')
                                                paragraph.add_run(bold_text).bold = True
                                            elif part.startswith('`') and part.endswith('`'):
                                                # Code/inline code
                                                code_text = part.replace('`', '')
                                                run = paragraph.add_run(code_text)
                                                run.font.name = 'Courier New'
                                            else:
                                                # Regular text
                                                if part:
                                                    paragraph.add_run(part)
                                    
                                    docx_doc = Document()
                                    docx_doc.add_heading(f'Summary: {doc["filename"]}', 0)
                                    
                                    # Parse markdown and add to document with proper formatting
                                    lines = detailed_summary.split('\n')
                                    for line in lines:
                                        line = line.strip()
                                        if not line:
                                            continue
                                        
                                        # Check if it's a heading (starts and ends with **)
                                        if re.match(r'^\*\*.*\*\*$', line):
                                            # Remove all ** and use as heading
                                            heading_text = re.sub(r'\*\*', '', line).strip()
                                            docx_doc.add_heading(heading_text, level=2)
                                        elif line.startswith('•') or line.startswith('-') or (line.startswith('*') and not line.startswith('**')):
                                            # Bullet point - remove bullet and process inline formatting
                                            bullet_text = re.sub(r'^[•\-\*]\s*', '', line)
                                            p = docx_doc.add_paragraph(style='List Bullet')
                                            add_formatted_text(p, bullet_text)
                                        else:
                                            # Regular paragraph - process inline formatting
                                            p = docx_doc.add_paragraph()
                                            add_formatted_text(p, line)
                                    
                                    # Save to BytesIO
                                    docx_buffer = BytesIO()
                                    docx_doc.save(docx_buffer)
                                    docx_buffer.seek(0)
                                    
                                    st.download_button(
                                        label="📝 Download as DOCX",
                                        data=docx_buffer.getvalue(),
                                        file_name=f"{doc['filename']}_summary.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"download_docx_{doc['filename']}"
                                    )
                                except Exception as e:
                                    st.info(f"DOCX unavailable")
                        else:
                            st.info("Detailed summary is being generated...")
                
                with col2:
                    if is_lead:
                        if st.button("🗑️ Delete", key=f"delete_{doc['filename']}", type="secondary"):
                            if document_store.delete_document(doc['filename'], team_id):
                                st.success(f"Document '{doc['filename']}' deleted successfully!")
                                st.rerun()
                            else:
                                st.error("Failed to delete document.")

    # Tab 3: Chatbot
    with tab3:
        # Initialize chat history in session state
        if "messages" not in st.session_state:
            st.session_state.messages = []
        
        # Initialize pending prompt flag
        if "pending_prompt" not in st.session_state:
            st.session_state.pending_prompt = None

        # Create a container for messages (allows scrolling)
        messages_container = st.container()
        
        with messages_container:
            # Show welcome message if no messages exist
            if len(st.session_state.messages) == 0 and not st.session_state.pending_prompt:
                st.markdown("""
                <div style='display: flex; justify-content: center; align-items: center; min-height: 400px; text-align: center;'>
                    <div>
                        <p style='color: #64748b; font-weight: 400; font-size: 2.5rem; margin-bottom: 1.5rem;'>Welcome to your Knowledge Base</p>
                        <p style='color: #94a3b8; font-size: 1.1rem;'>Ask any questions about your project documents, and I'll help you find the answers.</p>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Display all chat messages from history (renders above input)
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
            
            # Check if there's a pending prompt to process
            if st.session_state.pending_prompt:
                # Show thinking indicator below the last message
                with st.chat_message("assistant"):
                    with st.spinner("Generating response..."):
                        team_id = st.session_state.get("team_id")
                        team_name = st.session_state.get("team_name")
                        # Use Pinecone vector search for relevant context
                        context = document_store.get_context_for_query(st.session_state.pending_prompt, team_id, team_name)
                        response = gemini_utils.chat_with_documents(st.session_state.pending_prompt, context)
                        st.markdown(response)
                
                # Add assistant message to history
                st.session_state.messages.append({"role": "assistant", "content": response})
                
                # Clear pending prompt
                st.session_state.pending_prompt = None
                
                # Rerun to refresh the UI
                st.rerun()

        # Chat Input (automatically pinned to bottom when in main body)
        if prompt := st.chat_input("Ask any questions related to your project here"):
            # Add user message to history immediately
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            # Set pending prompt to trigger response generation on next rerun
            st.session_state.pending_prompt = prompt
            
            # Rerun immediately so user message appears above input
            st.rerun()
    
    # Tab 4: Team Settings (only for team leads)
    if is_lead:
        with tab4:
            st.header("⚙️ Team Settings")
            st.info("As the Team Lead, you can manage team settings here.")
            
            # Get current team info
            team = team_store.get_team_by_id(team_id) if team_id else None
            
            if team:
                st.markdown("### Current Team Information")
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Team Name:** {team.get('name', 'N/A')}")
                    st.write(f"**Access Code:** {team.get('access_code', 'N/A')}")
                with col2:
                    current_lead = team.get('team_lead_email', 'Not assigned')
                    st.write(f"**Current Team Lead:** {current_lead}")
                    team_created_formatted = format_date_ist(team.get('created_at', 'N/A'))
                    st.write(f"**Team Created:** {team_created_formatted}")
                
                st.markdown("---")
                st.markdown("### Update Team Lead")
                st.caption("Assign a new team lead by entering their email address.")
                
                with st.form("update_team_lead_form"):
                    new_lead_email = st.text_input(
                        "New Team Lead Email",
                        placeholder="Enter the email of the new team lead",
                        help="The new team lead must be a member of this team"
                    )
                    submit_update = st.form_submit_button("Update Team Lead", type="primary")
                    
                    if submit_update:
                        if not new_lead_email:
                            st.error("Please enter an email address.")
                        elif not user_store.validate_email(new_lead_email):
                            st.error("Please enter a valid email address.")
                        else:
                            # Check if the email belongs to a user in this team
                            new_lead_user = user_store.get_user_by_email(new_lead_email)
                            if not new_lead_user:
                                st.error(f"No user found with email '{new_lead_email}'. They must sign up first.")
                            elif new_lead_user.get("team_id") != team_id:
                                st.error(f"This user belongs to a different team.")
                            else:
                                # Update team lead
                                if team_store.update_team_lead(team_id, new_lead_email):
                                    st.success(f"Team lead updated successfully! New team lead: {new_lead_email}")
                                    st.info("The new team lead will have full access after they log out and log back in.")
                                    # Update session state if updating to current user
                                    if new_lead_email.lower() == user_email.lower():
                                        st.session_state.is_team_lead = True
                                    else:
                                        st.session_state.is_team_lead = False
                                        st.warning("You are no longer the team lead. Please refresh the page.")
                                    st.rerun()
                                else:
                                    st.error("Failed to update team lead.")
            else:
                st.error("Team information not found.")
