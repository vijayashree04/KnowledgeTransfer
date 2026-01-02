import os
from datetime import datetime
from io import BytesIO
from typing import Optional, List, Dict
from supabase_config import supabase
import vector_store
from pypdf import PdfReader
import docx
from uuid import UUID

def _check_supabase():
    """Checks if Supabase is configured."""
    if supabase is None:
        raise ValueError("Supabase is not configured. Please set SUPABASE_URL and SUPABASE_KEY environment variables.")

def _validate_team_id(team_id: str) -> str:
    """Validates and converts team_id to UUID string."""
    if not team_id:
        raise ValueError("team_id is required.")
    try:
        return str(UUID(team_id))
    except ValueError:
        raise ValueError(f"Invalid team_id format: {team_id}. Expected UUID.")

def save_uploaded_file(uploaded_file, username: str, team_id: Optional[str] = None, team_name: Optional[str] = None) -> Dict:
    """Saves an uploaded file to Supabase Storage and updates metadata in Supabase."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to save documents.")
    
    # Read file content first (before uploading)
    file_bytes = uploaded_file.getbuffer()
    content = ""
    try:
        name_lower = uploaded_file.name.lower()
        raw_bytes = file_bytes.tobytes()

        # Simple text-based formats
        if name_lower.endswith((".txt", ".md", ".py", ".js", ".json")):
            content = raw_bytes.decode("utf-8", errors="ignore")

        # PDF parsing
        elif name_lower.endswith(".pdf"):
            try:
                reader = PdfReader(BytesIO(raw_bytes))
                pages_text = []
                for page in reader.pages:
                    try:
                        pages_text.append(page.extract_text() or "")
                    except Exception:
                        continue
                content = "\n".join(pages_text).strip() or "[No extractable text found in PDF]"
            except Exception as e_pdf:
                content = f"[Error reading PDF content: {e_pdf}]"

        # Word document (.docx)
        elif name_lower.endswith(".docx"):
            try:
                doc = docx.Document(BytesIO(raw_bytes))
                paras = [p.text for p in doc.paragraphs if p.text]
                content = "\n".join(paras).strip() or "[No extractable text found in DOCX]"
            except Exception as e_docx:
                content = f"[Error reading DOCX content: {e_docx}]"

        # Legacy .doc: best-effort text decode
        elif name_lower.endswith(".doc"):
            try:
                content = raw_bytes.decode("utf-8", errors="ignore")
            except Exception:
                content = "[Binary .doc file - text extraction not implemented]"

        # Fallback: best-effort text decode
        else:
            try:
                content = raw_bytes.decode("utf-8", errors="ignore")
            except Exception:
                content = "[Binary file content - Text extraction required]"
    except Exception as e:
        content = f"Error reading file: {str(e)}"

    # Validate team_id
    team_uuid = _validate_team_id(team_id)
    
    # Upload to Supabase Storage
    storage_path = f"{team_uuid}/{uploaded_file.name}"
    storage_bucket = "documents"  # Default bucket name
    
    try:
        # Upload file to Supabase Storage
        storage_response = supabase.storage.from_(storage_bucket).upload(
            storage_path,
            file_bytes.tobytes(),
            file_options={"content-type": uploaded_file.type or "text/plain", "upsert": "true"}
        )
        storage_url = None
        # Get public URL if bucket is public
        try:
            storage_url = supabase.storage.from_(storage_bucket).get_public_url(storage_path)
        except:
            # If bucket is private, URL will be None
            pass
        
        # Use storage path instead of local file path
        file_path = f"storage://{storage_bucket}/{storage_path}"
        
    except Exception as storage_error:
        raise ValueError(f"Failed to upload file to Supabase Storage: {str(storage_error)}")
    
    # Check if document already exists (same filename + team_id)
    existing = supabase.table("documents").select("id").eq("filename", uploaded_file.name).eq("team_id", team_uuid).execute()
    
    doc_data = {
        "filename": uploaded_file.name,
        "uploaded_by": username,
        "team_id": team_uuid,
        "file_path": file_path,
        "summary": None,
        "content": content
    }
    
    if existing.data and len(existing.data) > 0:
        # Update existing document
        result = supabase.table("documents").update(doc_data).eq("id", existing.data[0]["id"]).execute()
    else:
        # Insert new document
        result = supabase.table("documents").insert(doc_data).execute()
    
    if result.data and len(result.data) > 0:
        doc = result.data[0]
        metadata = {
            "filename": uploaded_file.name,
            "uploaded_by": username,
            "upload_date": doc.get("upload_date", datetime.now().isoformat()),
            "path": file_path,
            "summary": None,
            "content": content,
            "team_id": team_id,
            "team_name": team_name,
            "id": doc.get("id"),
        }

        # Index document in Pinecone for RAG (best-effort, non-blocking)
        try:
            if vector_store.is_enabled():
                vector_store.index_document(
                    team_id=str(team_id),
                    doc_id=str(metadata.get("id")),
                    filename=metadata["filename"],
                    content=metadata["content"],
                    uploaded_by=metadata["uploaded_by"],
                    summary=metadata.get("summary"),
                    team_name=metadata.get("team_name"),
                )
        except Exception as e:
            print(f"Warning: failed to index document in Pinecone: {e}")

        return metadata
    else:
        raise ValueError("Failed to save document metadata to Supabase.")

def get_documents(team_id: Optional[str] = None) -> List[Dict]:
    """Retrieves all document metadata for a team from Supabase."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to retrieve documents.")
    
    try:
        team_uuid = _validate_team_id(team_id)
        
        result = supabase.table("documents").select("*").eq("team_id", team_uuid).order("upload_date", desc=True).execute()
        
        if result.data:
            return [_format_document(doc) for doc in result.data]
        return []
        
    except Exception as e:
        raise ValueError(f"Failed to get documents: {str(e)}")

def update_document_summary(filename: str, summary: str, team_id: Optional[str] = None) -> bool:
    """Updates the summary for a specific document in Supabase."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to update document summary.")
    
    try:
        team_uuid = _validate_team_id(team_id)
        
        # Find document by filename and team_id
        result = supabase.table("documents").select("id").eq("filename", filename).eq("team_id", team_uuid).execute()
        
        if result.data and len(result.data) > 0:
            doc_id = result.data[0]["id"]
            update_result = supabase.table("documents").update({
                "summary": summary
            }).eq("id", doc_id).execute()
            
            return update_result.data is not None and len(update_result.data) > 0
        
        return False
        
    except Exception as e:
        raise ValueError(f"Failed to update document summary: {str(e)}")

def save_summary_file_to_storage(filename: str, summary_text: str, team_id: Optional[str] = None, file_format: str = "docx") -> Optional[str]:
    """Saves the summarized document as a file to Supabase Storage and returns the storage path."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to save summary file.")
    
    try:
        team_uuid = _validate_team_id(team_id)
        storage_bucket = "documents"
        
        # Create summary filename
        base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
        summary_filename = f"{base_name}_summary.{file_format}"
        storage_path = f"{team_uuid}/summaries/{summary_filename}"
        
        # Prepare file content based on format
        if file_format == "docx":
            # Create Word document
            from docx import Document
            from io import BytesIO
            import re
            
            docx_doc = Document()
            docx_doc.add_heading(f'Summary: {filename}', 0)
            
            # Helper function to add text with bold formatting
            def add_formatted_text(paragraph, text):
                """Adds text to paragraph, handling **bold** and `code` markdown."""
                # Split by ** and ` to find formatted sections
                parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        bold_text = part.replace('**', '')
                        paragraph.add_run(bold_text).bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Code/inline code
                        code_text = part.replace('`', '')
                        run = paragraph.add_run(code_text)
                        run.font.name = 'Courier New'
                    else:
                        # Regular text
                        if part:
                            paragraph.add_run(part)
            
            # Parse markdown and add to document
            lines = summary_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if it's a heading (starts and ends with **)
                if re.match(r'^\*\*.*\*\*$', line):
                    # Remove all ** and use as heading
                    heading_text = re.sub(r'\*\*', '', line).strip()
                    docx_doc.add_heading(heading_text, level=2)
                elif line.startswith('•') or line.startswith('-') or (line.startswith('*') and not line.startswith('**')):
                    # Bullet point - remove bullet and process inline formatting
                    bullet_text = re.sub(r'^[•\-\*]\s*', '', line)
                    p = docx_doc.add_paragraph(style='List Bullet')
                    add_formatted_text(p, bullet_text)
                else:
                    # Regular paragraph - process inline formatting
                    p = docx_doc.add_paragraph()
                    add_formatted_text(p, line)
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            docx_doc.save(docx_buffer)
            file_content = docx_buffer.getvalue()
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif file_format == "txt":
            file_content = summary_text.encode('utf-8')
            content_type = "text/plain"
        elif file_format == "md":
            file_content = summary_text.encode('utf-8')
            content_type = "text/markdown"
        else:
            # Default to docx
            file_content = summary_text.encode('utf-8')
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Upload to Supabase Storage
        supabase.storage.from_(storage_bucket).upload(
            storage_path,
            file_content,
            file_options={"content-type": content_type, "upsert": "true"}
        )
        
        # Return storage path
        return f"storage://{storage_bucket}/{storage_path}"
        
    except Exception as e:
        print(f"Warning: Failed to save summary file to storage: {e}")
        return None

def update_document_detailed_summary(filename: str, detailed_summary: str, team_id: Optional[str] = None) -> bool:
    """Updates the detailed summary for a specific document in Supabase and saves summary file to storage."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to update document detailed summary.")
    
    try:
        team_uuid = _validate_team_id(team_id)
        
        # Find document by filename and team_id
        result = supabase.table("documents").select("id").eq("filename", filename).eq("team_id", team_uuid).execute()
        
        if result.data and len(result.data) > 0:
            doc_id = result.data[0]["id"]
            
            # Save summary file to storage as Word document
            summary_file_path = save_summary_file_to_storage(filename, detailed_summary, team_id, "docx")
            
            # Prepare update data
            update_data = {
                "detailed_summary": detailed_summary
            }
            
            # Add summary_file_path if file was saved successfully
            if summary_file_path:
                update_data["summary_file_path"] = summary_file_path
            
            # Update detailed_summary and summary_file_path columns
            update_result = supabase.table("documents").update(update_data).eq("id", doc_id).execute()
            
            return update_result.data is not None and len(update_result.data) > 0
        
        return False
        
    except Exception as e:
        # If column doesn't exist, this will fail - user needs to run migration
        raise ValueError(f"Failed to update document detailed summary: {str(e)}. Please run the migration to add 'detailed_summary' and 'summary_file_path' columns to documents table.")

def get_all_context(team_id: Optional[str] = None) -> str:
    """Combines all document content for the chatbot context from Supabase."""
    docs = get_documents(team_id)
    context = ""
    for doc in docs:
        context += f"\n\n--- Document: {doc['filename']} ---\n"
        context += doc.get("content", "")
        if doc.get("summary"):
            context += f"\nSummary: {doc['summary']}\n"
    return context

def get_context_for_query(query: str, team_id: Optional[str] = None, team_name: Optional[str] = None) -> str:
    """Gets relevant context for a query using Pinecone vector search, with fallback to all documents."""
    # Try to use Pinecone vector search if available
    if vector_store.is_enabled() and team_id:
        try:
            # Query Pinecone for similar documents
            matches = vector_store.query_similar_documents(
                query_text=query,
                team_id=str(team_id),
                team_name=team_name,
                top_k=10  # Get top 10 most relevant chunks (allows multiple chunks per document)
            )
            
            if matches:
                # Build context from Pinecone matches
                context = vector_store.build_context_from_matches(matches)
                if context:
                    return context
        except Exception as e:
            print(f"Warning: Pinecone query failed, falling back to all documents: {e}")
    
    # Fallback to all documents if Pinecone is not available or query failed
    return get_all_context(team_id)

def delete_document(filename: str, team_id: Optional[str] = None) -> bool:
    """Deletes a document from Supabase."""
    _check_supabase()
    
    if not team_id:
        raise ValueError("team_id is required to delete documents.")
    
    try:
        team_uuid = _validate_team_id(team_id)
        
        # Find and delete document
        result = supabase.table("documents").select("id, file_path, summary_file_path").eq("filename", filename).eq("team_id", team_uuid).execute()
        
        if result.data and len(result.data) > 0:
            doc_id = result.data[0]["id"]
            file_path = result.data[0].get("file_path")
            summary_file_path = result.data[0].get("summary_file_path")
            
            # Delete original file from storage
            if file_path and file_path.startswith("storage://"):
                # Delete from Supabase Storage
                try:
                    # Extract bucket and path from storage://bucket/path format
                    storage_path = file_path.replace("storage://documents/", "")
                    supabase.storage.from_("documents").remove([storage_path])
                except Exception as e:
                    print(f"Warning: Could not delete original file from storage: {e}")
            
            # Delete summary file from storage
            if summary_file_path and summary_file_path.startswith("storage://"):
                try:
                    # Extract bucket and path from storage://bucket/path format
                    summary_storage_path = summary_file_path.replace("storage://documents/", "")
                    supabase.storage.from_("documents").remove([summary_storage_path])
                except Exception as e:
                    print(f"Warning: Could not delete summary file from storage: {e}")
            
            # Delete from database
            delete_result = supabase.table("documents").delete().eq("id", doc_id).execute()
            
            # Delete vectors from Pinecone if enabled
            try:
                if vector_store.is_enabled():
                    vector_store.delete_document_vectors(
                        doc_id=str(doc_id),
                        team_id=str(team_uuid),
                        team_name=None
                    )
            except Exception as e:
                print(f"Warning: Could not delete vectors from Pinecone: {e}")
            
            return True
        
        return False
        
    except Exception as e:
        raise ValueError(f"Failed to delete document: {str(e)}")

def _format_document(doc: Dict) -> Dict:
    """Formats document data to match expected structure."""
    return {
        "id": doc.get("id"),
        "filename": doc.get("filename"),
        "uploaded_by": doc.get("uploaded_by"),
        "upload_date": doc.get("upload_date", datetime.now().isoformat()),
        "path": doc.get("file_path"),
        "summary": doc.get("summary"),
        "detailed_summary": doc.get("detailed_summary"),  # Add detailed_summary field
        "summary_file_path": doc.get("summary_file_path"),  # Add summary_file_path field
        "content": doc.get("content", ""),
        "team_id": doc.get("team_id")
    }
